from flask import Blueprint, request, jsonify, render_template, make_response
import time
from docx import Document
import tempfile
import os
import pandas as pd
from docx.shared import Inches
import io
from .preprocessing import (
    preprocess_puc_file, feature_engineering
)

from .predictions import set_flag_conditions as stp_conditions
from .tsx_predictions import set_flag_conditions as tsx_conditions

from .visualizations import (
    make_flagged,
    plot_sensor_values, plot_sensor_trends, 
    plot_door_histogram,
    plot_trend_issue_altair,
    plot_tc10, plot_tc1_tc6
)
from .summary import generate_summary

# imports for file download
import re
import base64

# For Word export
main = Blueprint('main', __name__)

DF = None
CHARTS: dict = {}
FLAGGED = None

@main.route('/')
def index():
    return render_template('index.html')

@main.route('/process', methods=['POST'])
def process_file():
    file = request.files['file']
    raw_data = file.read()
    
    package = preprocess_puc_file(raw_data)
    
    if package is None:
        return jsonify({"status": "error", "message": "The data is lesser than 45 days for analysis more data requires, manual analysis required."}), 400
    
    df, door_events_original, power_events_df, ref_df, file_type, note = package
    
    
    df, tcs = feature_engineering(df)
    
    if 'Total Time of Opening (secs)' in door_events_original.columns:
        door_events_filtered = door_events_original[door_events_original['Total Time of Opening (secs)'] > 60]
        
    else:
        print("No 'Total Time of Opening (secs)' column found in door_events_original.")
        door_events_filtered = pd.DataFrame()  # Empty DataFrame
    
    if ('Date of Event' in power_events_df.columns) and ('Event' in power_events_df.columns):
        power_events_df = (
            power_events_df
            .groupby(['Date of Event', 'Event'])
            .size()
            .reset_index(name='Count of This Event on Day')
        )
        
    else:
        print("No 'Date of Event' and/or 'Event' columns found in power_events_df.")
        power_events_df = pd.DataFrame()  # Empty DataFrame

    if file_type == 'TSX':
        df = tsx_conditions(df, ref_df)
        
    else:
        df = stp_conditions(df)

    filtered = df.loc[df['Sustained_Issue'] == True, ['Date/Time', 'Trend_Flag']]
    
    summary= generate_summary(df, door_events_filtered, power_events_df, door_events_original, tcs, ref_df, filtered)
    
    summary['file_type'] = file_type
    summary['note'] = note
    
    # global variables for use across routes
    global DF, FLAGGED, CHARTS
    FLAGGED = make_flagged(filtered)
    DF = df
    
    door_events_chart = plot_door_histogram(door_events_original) 
    CHARTS['Door Events'] = door_events_chart
    
    return jsonify(summary)

@main.route('/visualizations', methods=['POST', 'GET'])
def visualizations():
    global DF, FLAGGED, CHARTS
    
    if DF is None:
        return jsonify({"error": "No data available for visualizations"}), 400

    try:
        DF['Date/Time'] = pd.to_datetime(DF['Date/Time'])
    except Exception:
        return jsonify({"error": "Invalid or missing 'Date/Time' column"}), 400

    # Build safe sensor_values_df
    requested_cols = ['Date/Time', 'RTD','Setpoint', 'TC1','TC2', 'TC10', 'TC3','TC8', 'TC4', 'TC6']
    available_cols = [col for col in requested_cols if col in DF.columns]
    sensor_values_df = DF[available_cols].dropna(axis=1, how='all')

    # Build safe sensor_trends_df
    requested_trend_cols = ['Date/Time', 'RTD_trend', 'Stage 1 RPM', 'Stage 2 RPM',
                            'TC1_trend','TC2_trend', 'TC10_trend','TC8_trend',
                            'TC3_trend', 'TC4_trend', 'TC6_trend', 'Trend_Flag']
    available_trend_cols = [col for col in requested_trend_cols if col in DF.columns]
    sensor_trends_df = DF[available_trend_cols].dropna(axis=1, how='all')

    response = {}

    # Only add flagged charts if FLAGGED is valid
    if FLAGGED is not None and not FLAGGED.empty:
        try:
            response["sensor_values"] = plot_sensor_values(sensor_values_df, FLAGGED)  # type: ignore
            response["sensor_trends"] = plot_sensor_trends(sensor_trends_df, FLAGGED)  # type: ignore
            CHARTS['Sensor Values'] = response["sensor_values"]
            CHARTS['Sensor Trends'] = response["sensor_trends"]
        except Exception as e:
            import logging; logging.exception("Failed to generate FLAGGED charts")

    # Other charts
    try:
        response["trend_issues_altair"] = plot_trend_issue_altair(DF).to_dict(format='vega') #type: ignore
    except Exception as e:
        response["trend_issues_altair"] = {}
        import logging; logging.exception("Failed to generate trend issues chart")

    try:
        response["tc10_chart"] = plot_tc10(DF).to_dict(format='vega') #type: ignore
    except Exception as e:
        response["tc10_chart"] = {}
        import logging; logging.exception("Failed to generate TC10 chart")

    try:
        response["tc1_tc6_chart"] = plot_tc1_tc6(DF).to_dict(format='vega')
    except Exception as e:
        response["tc1_tc6_chart"] = {}
        import logging; logging.exception("Failed to generate TC1-TC6 chart")

    # Door Events handled safely
    door_events_chart = CHARTS.get('Door Events')
    if door_events_chart is not None and hasattr(door_events_chart, 'to_dict'):
        response["door events"] = door_events_chart.to_dict(format='vega')
    else:
        response["door events"] = {}

    return jsonify(response)

@main.route('/download_word', methods=['POST'])
def download_word():
    """Generate and send a Word document with the summary."""
    print("Downloading.......")

    start = time.perf_counter()

    data = request.get_json()
    doc = Document()
    doc.add_heading(data.get('title', 'Telemetry Summary'), 0)

    # 1. Add Observations first
    observation = data.get('observation')
    if observation:
        doc.add_heading('Observation', level=1)
        if isinstance(observation, str):
            doc.add_paragraph(observation)
        elif isinstance(observation, list):
            for item in observation:
                doc.add_paragraph(str(item))
        elif isinstance(observation, dict):
            for k, v in observation.items():
                doc.add_paragraph(f"{k}: {v}")
        else:
            doc.add_paragraph(str(observation))

    # 2. Add charts immediately after Observations
    with tempfile.TemporaryDirectory() as tmpdir:
        for chart_title, chart_obj in CHARTS.items():
            try:
                safe_name = chart_title.replace(" ", "_") + ".png"
                chart_path = os.path.join(tmpdir, safe_name)

                # Altair chart with .save
                if hasattr(chart_obj, "save") and callable(chart_obj.save):
                    chart_obj.save(chart_path, format='png')

                # Matplotlib (base64-encoded PNG string)
                elif isinstance(chart_obj, str) and chart_obj.startswith("data:image/png;base64,"):
                    base64_data = re.sub('^data:image/.+;base64,', '', chart_obj)
                    with open(chart_path, "wb") as f:
                        f.write(base64.b64decode(base64_data))

                else:
                    raise TypeError(f"Unknown chart type for '{chart_title}'")

                doc.add_heading(chart_title, level=1)
                doc.add_picture(chart_path, width=Inches(6))

            except Exception as e:
                doc.add_paragraph(f"Failed to insert chart '{chart_title}': {str(e)}")

    # 3. Add remaining sections
    for key in ['Summary: Events', '🧠 Root Cause Explanation:']:
        value = data.get(key)
        if value:
            doc.add_heading(key.replace('_', ' ').title(), level=1)
            if isinstance(value, str):
                doc.add_paragraph(value)
            elif isinstance(value, list):
                for item in value:
                    doc.add_paragraph(str(item))
            elif isinstance(value, dict):
                for k, v in value.items():
                    doc.add_paragraph(f"{k}: {v}")
            else:
                doc.add_paragraph(str(value))

    # Generate response
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    response = make_response(buf.read())
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    response.headers['Content-Disposition'] = f'attachment; filename=telemetry_summary_{time.strftime("%Y%m%d")}.docx'

    end = time.perf_counter()
    print(f"Total download time = {end - start}")

    return response